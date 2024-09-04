import os

from docx.oxml import OxmlElement, ns
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import docx
from PyQt5.QtWidgets import QMessageBox
from PyQt5 import QtGui
import webbrowser
from pdf2docx import Converter
import win32com.client
import re

path = f'{os.environ["USERPROFILE"]}\\Desktop\\Template-uri'

def makefolder():
    isExist = os.path.exists(path)
    if not isExist:
        os.makedirs(path)

def succesMessage():
    msg = QMessageBox()
    msg.setWindowIcon(QtGui.QIcon(":/imgs/favicon.ico"))
    msg.setIcon(QMessageBox.Information)
    msg.setText("Felicitări ! \nTemplate-ul dvs. a fost creat cu succes !")
    msg.setWindowTitle("Succes !!")
    msg.exec_()


def teza(dates, date):
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldStart = create_element('w:fldChar')
        create_attribute(fldStart, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'separate')

        fldChar2 = create_element('w:t')
        fldChar2.text = "2"

        fldEnd = create_element('w:fldChar')
        create_attribute(fldEnd, 'w:fldCharType', 'end')

        run._r.append(fldStart)

        run._r.append(instrText)
        run._r.append(fldChar1)
        run._r.append(fldChar2)

        run._r.append(fldEnd)

    doc = docx.Document()
    add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.59)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)

    styles = doc.styles
    charstyle = styles.add_style('TNR_12', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    styles = doc.styles
    charstyle = styles.add_style('TNR_11', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Times New Roman'

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart.paragraph_format.space_after = Pt(0)
    standart.add_run('MINISTERUL EDUCAŢIEI, CULTURII ȘI CERCETĂRII AL REPUBLICII MOLDOVA', style='TNR_14')

    antet = doc.add_paragraph()
    antet.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    antet.paragraph_format.space_after = Pt(0)
    antet.add_run('Universitatea de Stat din Tiraspol', style='TNR_14')

    antet1 = doc.add_paragraph()
    antet1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # antet1.paragraph_format.line_spacing = Inches(0.2)
    antet1.paragraph_format.space_after = Pt(0)
    antet1.add_run('Facultatea Fizică Matematică şi Tehnologii Informaţionale', style='TNR_14')

    catedra = doc.add_paragraph()
    catedra.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    catedra.paragraph_format.line_spacing = Inches(0.2)
    catedra.add_run(dates[1].upper(), style='TNR_14').bold = True

    doc.add_paragraph().add_run()

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.paragraph_format.line_spacing = Inches(0.2)
    standart1.add_run(f'TEZĂ DE {date}'.upper(), style='TNR_14').bold = True

    doc.add_paragraph().add_run()
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.paragraph_format.line_spacing = Inches(0.2)
    title.add_run(dates[2].title(), style='TNR_14').bold = True

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    specialitate = doc.add_paragraph()
    specialitate.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # specialitate.paragraph_format.space_before = Pt(110)
    # specialitate.paragraph_format.space_after = Pt(1)
    specialitate.add_run(f'Specialitate: {dates[3].title()}', style='TNR_12').bold = True

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    table.cell(0, 0).width = Inches(1.96)
    table.cell(0, 1).width = Inches(1.96)
    table.cell(1, 0).width = Inches(1.96)
    table.cell(1, 1).width = Inches(1.96)
    table.cell(2, 0).width = Inches(1.96)
    table.cell(2, 1).width = Inches(1.96)
    table.cell(3, 0).width = Inches(1.96)
    table.cell(3, 1).width = Inches(1.96)
    table.cell(4, 0).width = Inches(1.96)
    table.cell(4, 1).width = Inches(1.96)

    row = table.rows[0].cells
    unu = row[0].paragraphs[0]
    unu.paragraph_format.line_spacing = Inches(0.2)
    unu.paragraph_format.space_after = Pt(1)
    unu.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = unu.add_run('Autor:', style='TNR_12')
    run.bold = True

    row0 = table.rows[0].cells
    unu0 = row0[1].paragraphs[0]
    unu0.paragraph_format.line_spacing = Inches(0.2)
    unu0.paragraph_format.space_after = Pt(1)
    unu0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run0 = unu0.add_run('', style='TNR_12')
    run0.bold = True

    row1 = table.rows[1].cells
    unu1 = row1[0].paragraphs[0]
    unu1.paragraph_format.line_spacing = Inches(0.2)
    unu1.paragraph_format.space_after = Pt(1)
    unu1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu1.add_run(f'{dates[0].title()}, gr.{dates[5].upper()},\n frecvenţă (la) {dates[6].lower()}', style='TNR_12')
    unu2 = row1[1].paragraphs[0]
    unu2.paragraph_format.line_spacing = Inches(0.2)
    unu2.paragraph_format.space_after = Inches(0.3)
    unu2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run1 = unu2.add_run('_____________ (semnătura)', style='TNR_12')
    run1.bold = True

    row2 = table.rows[2].cells
    unu3 = row2[0].paragraphs[0]
    unu3.paragraph_format.line_spacing = Inches(0.2)
    unu3.paragraph_format.space_after = Pt(1)
    unu3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run3 = unu3.add_run('Conducător ştiinţific:', style='TNR_12')
    run3.bold = True
    unu31 = row2[0].paragraphs[0]
    unu31.paragraph_format.line_spacing = Inches(0.2)
    unu31.paragraph_format.space_after = Pt(1)
    unu31.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu31.add_run(f'\n{dates[4].title()}, ', style='TNR_12')
    unu32 = row2[0].paragraphs[0]
    unu32.paragraph_format.line_spacing = Inches(0.2)
    unu32.paragraph_format.space_after = Inches(0.3)
    unu32.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu32.add_run('\ndr. conf. univ., ', style='TNR_12')

    unu4 = row2[1].paragraphs[0]
    unu4.paragraph_format.line_spacing = Inches(0.2)
    unu4.paragraph_format.space_after = Pt(1)
    unu4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run4 = unu4.add_run('_____________ (semnătura)', style='TNR_12')
    run4.bold = True

    row3 = table.rows[3].cells
    unu5 = row3[0].paragraphs[0]
    unu5.paragraph_format.line_spacing = Inches(0.2)
    unu5.paragraph_format.space_after = Pt(1)
    unu5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run5 = unu5.add_run('Admis la susţinere \n', style='TNR_12')
    run5.bold = True
    unu51 = row3[0].paragraphs[0]
    unu51.paragraph_format.line_spacing = Inches(0.2)
    unu51.paragraph_format.space_after = Pt(1)
    unu51.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu5.add_run('Şef Catedra ITI,', style='TNR_12')
    unu6 = row3[1].paragraphs[0]
    unu6.paragraph_format.line_spacing = Inches(0.2)
    unu6.paragraph_format.space_after = Inches(0.3)
    unu6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run6 = unu6.add_run('_____________ (data, semnătura)', style='TNR_12')
    run6.bold = True

    row4 = table.rows[4].cells
    unu7 = row4[0].paragraphs[0]
    unu7.paragraph_format.line_spacing = Inches(0.2)
    unu7.paragraph_format.space_after = Pt(1)
    unu7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu7.add_run('Liubomir Chiriac \ndr. habilitat, prov. univ ', style='TNR_12')
    unu8 = row4[1].paragraphs[0]
    unu8.paragraph_format.line_spacing = Inches(0.2)
    unu8.paragraph_format.space_after = Pt(1)
    unu8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run8 = unu8.add_run('', style='TNR_12')
    run8.bold = True

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    ch = doc.add_paragraph()
    ch.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ch.paragraph_format.space_before = Pt(100)
    ch.add_run('Chișinău, 2022', style='TNR_14')

    doc.add_paragraph().add_run()

    title1 = doc.add_paragraph()
    title1.paragraph_format.space_after = Pt(0)
    title1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title1.add_run('CUPRINS:', style='TNR_14').bold = True
    doc.add_paragraph().add_run()
    title2 = doc.add_paragraph()
    # title2.paragraph_format.space_after = Pt(0)
    title2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title2.add_run(
        'ADNOTARE..............................................................................................................2',
        style='TNR_14').bold = True

    title3 = doc.add_paragraph()
    # title3.paragraph_format.space_after = Pt(0)
    title3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title3.add_run(
        'LISTA ABREVIERILOR...........................................................................................3',
        style='TNR_14').bold = True

    title4 = doc.add_paragraph()
    # title4.paragraph_format.space_after = Pt(0)
    title4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title4.add_run(
        'Introducere..................................................................................................................4',
        style='TNR_14').bold = True

    title5 = doc.add_paragraph()
    title5.paragraph_format.space_after = Pt(0)
    title5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title5.add_run(
        'CAPITOLUL I. DENUMIREA CAPITOLULUI.....................................................7',
        style='TNR_14').bold = True

    title6 = doc.add_paragraph()
    title6.paragraph_format.space_after = Pt(0)
    title6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title6.add_run(
        '1.1. Denumirea subcapitolului.......................................................................................7',
        style='TNR_14')

    title7 = doc.add_paragraph()
    title7.paragraph_format.space_after = Pt(0)
    title7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title7.add_run(
        '1.2. Denumirea subcapitolului.......................................................................................9',
        style='TNR_14')

    title8 = doc.add_paragraph()
    # title8.paragraph_format.space_after = Pt(0)
    title8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title8.add_run(
        '1.3. Denumirea subcapitolului......................................................................................13',
        style='TNR_14')

    title9 = doc.add_paragraph()
    title9.paragraph_format.space_after = Pt(0)
    title9.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title9.add_run(
        'CAPITOLUL II. DENUMIREA CAPITOLULUI.................................................18',
        style='TNR_14').bold = True

    title10 = doc.add_paragraph()
    title10.paragraph_format.space_after = Pt(0)
    title10.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title10.add_run(
        '2.1. Denumirea subcapitolului.....................................................................................18',
        style='TNR_14')

    title11 = doc.add_paragraph()
    title11.paragraph_format.space_after = Pt(0)
    title11.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title11.add_run(
        '2.2. Denumirea subcapitolului.....................................................................................20',
        style='TNR_14')

    title12 = doc.add_paragraph()
    # title12.paragraph_format.space_after = Pt(0)
    title12.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title12.add_run(
        '2.3. Denumirea subcapitolului.....................................................................................25',
        style='TNR_14')

    title14 = doc.add_paragraph()
    title14.paragraph_format.space_after = Pt(0)
    title14.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title14.add_run(
        'CAPITOLUL III. DENUMIREA CAPITOLULUI................................................28',
        style='TNR_14').bold = True

    title15 = doc.add_paragraph()
    title15.paragraph_format.space_after = Pt(0)
    title15.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title15.add_run(
        '3.1. Denumirea subcapitolului......................................................................................28',
        style='TNR_14')

    title16 = doc.add_paragraph()
    title16.paragraph_format.space_after = Pt(0)
    title16.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title16.add_run(
        '3.2. Denumirea subcapitolului......................................................................................29',
        style='TNR_14')

    title17 = doc.add_paragraph()
    # title17.paragraph_format.space_after = Pt(0)
    title17.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title17.add_run(
        '3.3. Denumirea subcapitolului......................................................................................30',
        style='TNR_14')

    title18 = doc.add_paragraph()
    # title18.paragraph_format.space_after = Pt(0)
    title18.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title18.add_run(
        'Concluzii și recomandări...........................................................................................31',
        style='TNR_14').bold = True

    title19 = doc.add_paragraph()
    # title19.paragraph_format.space_after = Pt(0)
    title19.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title19.add_run(
        'Bibliografie..................................................................................................................32',
        style='TNR_14').bold = True

    title20 = doc.add_paragraph()
    # title20.paragraph_format.space_after = Pt(0)
    title20.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title20.add_run(
        'ANEXE........................................................................................................................34',
        style='TNR_14').bold = True
    title21 = doc.add_paragraph()
    # title21.paragraph_format.space_after = Pt(0)
    title21.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title21.add_run(
        'Declaraţia privind asumarea răspunderii................................................................35',
        style='TNR_14').bold = True
    doc.add_page_break()

    title22 = doc.add_paragraph()
    title22.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title22.add_run('Adnotare [MODEL]'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title23 = doc.add_paragraph()
    title23.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title23.add_run(open('files/adnotare.txt', 'r', encoding='utf8').read(), style='TNR_11')

    title24 = doc.add_paragraph()
    title24.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title24.add_run('Lista Abrevierilor [MODEL]'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title25 = doc.add_paragraph()
    title25.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title25.add_run(open('files/abrevieri.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    title26 = doc.add_paragraph()
    title26.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title26.add_run('Introducere [Informație]'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title27 = doc.add_paragraph()
    title27.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title27.add_run(open('files/introducere.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    title28 = doc.add_paragraph()
    title28.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title28.add_run('Capitolul i [Informație]'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title29 = doc.add_paragraph()
    title29.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title29.add_run(open('files/capitol1.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title28m = doc.add_paragraph()
    title28m.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title28m.add_run('Ca model avem :'.upper(), style='TNR_14').bold = True

    doc.add_page_break()

    title28m1 = doc.add_paragraph()
    title28m1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title28m1.add_run('CAPITOLUL I. ASPECTE TEORETICE PRIVIND GESTIUNEA PROGRAMĂRILOR'.upper(),
                      style='TNR_12').bold = True

    title28m2 = doc.add_paragraph()
    title28m2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title28m2.add_run(
        '1.1. Baze tehnice și teoretice de realizare a sistemului de programare a unităților economice'.upper(),
        style='TNR_12').bold = True

    title29m = doc.add_paragraph()
    title29m.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title29m.add_run(open('files/capitol1mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title28m3 = doc.add_paragraph()
    title28m3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title28m3.add_run('1.2. Tipologia și particularitățile sistemelor de programare '.upper(),
                      style='TNR_12').bold = True

    title29m1 = doc.add_paragraph()
    title29m1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title29m1.add_run(open('files/capitol12mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title28m4 = doc.add_paragraph()
    title28m4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title28m4.add_run(
        '1.3. Rolul și importanța gestiunii sistemelor de programare. Justificarea soluțiilor de proiectare pentru automatizarea sarcinii'.upper(),
        style='TNR_12').bold = True

    title29m2 = doc.add_paragraph()
    title29m2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title29m2.add_run(open('files/capitol13mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    title30 = doc.add_paragraph()
    title30.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title30.add_run('Capitolul ii [Informație]'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title31 = doc.add_paragraph()
    title31.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title31.add_run(open('files/capitol2.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title31m = doc.add_paragraph()
    title31m.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title31m.add_run('Ca model avem :'.upper(), style='TNR_14').bold = True

    doc.add_page_break()

    title31m1 = doc.add_paragraph()
    title31m1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title31m1.add_run('CAPITOLUL II. ANALIZA SISTEMULUI DE PROGRAMARE A UNITĂȚII'.upper(), style='TNR_12').bold = True

    title31m2 = doc.add_paragraph()
    title31m2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title31m2.add_run('2.1. Determinarea cerinţelor sistemului'.upper(), style='TNR_12').bold = True

    title31m3 = doc.add_paragraph()
    title31m3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title31m3.add_run(open('files/capitol21mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title31m4 = doc.add_paragraph()
    title31m4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title31m4.add_run('2.2. Limbaje de programare'.upper(), style='TNR_12').bold = True

    title31m5 = doc.add_paragraph()
    title31m5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title31m5.add_run(open('files/capitol22mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title31m6 = doc.add_paragraph()
    title31m6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title31m6.add_run('2.3. Cerințe tehnice și operaționale pentru colectarea și procesarea datelor de intrare'.upper(),
                      style='TNR_12').bold = True

    title32m6 = doc.add_paragraph()
    title32m6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title32m6.add_run(open('files/capitol23mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    title32 = doc.add_paragraph()
    title32.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title32.add_run('Capitolul iii [Informație]'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title33 = doc.add_paragraph()
    title33.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title33.add_run(open('files/capitol3.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title41m = doc.add_paragraph()
    title41m.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title41m.add_run('Ca model avem :'.upper(), style='TNR_14').bold = True

    mod = doc.add_paragraph()
    mod.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mod.add_run('CAPITOLUL III. PROIECT DE GESTIUNE A PROGRAMĂRILOR'.upper(), style='TNR_12').bold = True

    title33m2 = doc.add_paragraph()
    title33m2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title33m2.add_run('3.1. Date de referință'.upper(), style='TNR_12').bold = True

    title33m3 = doc.add_paragraph()
    title33m3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title33m3.add_run(open('files/capitol31mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title33m4 = doc.add_paragraph()
    title33m4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title33m4.add_run('3.2. Descrierea modulelor'.upper(), style='TNR_12').bold = True

    title33m5 = doc.add_paragraph()
    title33m5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title33m5.add_run(open('files/capitol32mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    title33m6 = doc.add_paragraph()
    title33m6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title33m6.add_run('3.3. Cerințe de implementareși de testare a programului'.upper(),
                      style='TNR_12').bold = True

    title32m61 = doc.add_paragraph()
    title32m61.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title32m61.add_run(open('files/capitol33mod.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    conc = doc.add_paragraph()
    conc.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    conc.add_run('Concluzii şi recomandări'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    title33 = doc.add_paragraph()
    title33.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title33.add_run(open('files/concluzii.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    bib = doc.add_paragraph()
    bib.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    bib.add_run('Bibliografie'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    bib1 = doc.add_paragraph()
    bib1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    bib1.add_run(open('files/bibliografie.txt', 'r', encoding='utf8').read(), style='TNR_12')
    #    doc.add_page_break()
    doc.add_page_break()

    anex = doc.add_paragraph()
    anex.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    anex.add_run('Anexe'.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    anex1 = doc.add_paragraph()
    anex1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    anex1.add_run(open('files/anexe.txt', 'r', encoding='utf8').read(), style='TNR_12')

    doc.add_page_break()

    dec = doc.add_paragraph()
    dec.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dec.add_run('Declaraţia privind asumarea răspunderii. '.upper(), style='TNR_14').bold = True

    doc.add_paragraph()

    dec1 = doc.add_paragraph()
    dec1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dec1.add_run(open('files/declaratii.txt', 'r', encoding='utf8').read(), style='TNR_12')

    makefolder()
    doc.save(f'{path}\\Template_Teza_de_{date.title()}.docx')
    succesMessage()


def download_scribd_doc(link):
    if link.lower().startswith("https://"):
        scribd = link
        downloadLink = "https://compress-pdf.tacz.info/?fileurl=https://dl.downscribd.com/pdownload/" + scribd.split("/")[4] + "/" + scribd.split("/")[5] + "&title=" + "+".join(
            scribd.split("/")[5].split("-")) + "&utm_source=downscr&utm_medium=queue&utm_campaign=dl"
        succesMessage_scribd()
        webbrowser.open(downloadLink)
    else:
        criticMessage()


def succesMessage_scribd():
    msg = QMessageBox()
    msg.setWindowIcon(QtGui.QIcon(":/imgs/favicon.ico"))
    msg.setIcon(QMessageBox.Information)
    msg.setText("Tastați OK pentru a descărca documentul !")
    msg.setWindowTitle("Felicitări !")
    msg.exec_()


def convert_pdf(pdf_file, docx_file):
    if (re.search("^[a-zA-Z]", pdf_file) or re.search("^[a-zA-Z]", docx_file)) is not None:
        cv = Converter(pdf_file)
        cv.convert(docx_file)
        cv.close()
        succesMessage_pdf()
    else:
        criticMessage()


def succesMessage_pdf():
    msg = QMessageBox()
    msg.setWindowIcon(QtGui.QIcon(":/imgs/favicon.ico"))
    msg.setIcon(QMessageBox.Information)
    msg.setText("Documentul a fost convertit !")
    msg.setWindowTitle("Felicitări !")
    msg.exec_()

def convert(inp, outp):
    file = open(outp, "w")
    file.close()
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(inp)
    doc.SaveAs(outp, FileFormat=17)
    doc.Close()
    word.Quit()

def convert_docx(docx_file, pdf_file):
    if (re.search("^[a-zA-Z]", pdf_file) or re.search("^[a-zA-Z]", docx_file)) is not None:
        convert(docx_file, pdf_file)
        succesMessage_pdf()
    else:
        criticMessage()


def criticMessage():
    msg = QMessageBox()
    msg.setWindowIcon(QtGui.QIcon(":/imgs/favicon.ico"))
    msg.setIcon(QMessageBox.Critical)
    msg.setText("Câmpuri goale ! Vă rugăm nu lăsați câmpuri goale !")
    msg.setWindowTitle("Eroare !")
    msg.exec_()


def contractPractica(date_cont):

    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.69)

    styles = doc.styles
    charstyle = styles.add_style('TNR_11', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Times New Roman'

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart.paragraph_format.space_after = Pt(0)
    standart.paragraph_format.line_spacing = Pt(0)
    standart.add_run('CONTRACT INDIVIDUAL DE REALIZARE A STAGIULUI DE PRACTICĂ', style='TNR_11').bold = True

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.paragraph_format.space_after = Pt(0)
    standart1.add_run('încheiat şi înregistrat cu nr. __ din ________', style='TNR_11')

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.paragraph_format.space_after = Pt(0)
    standart2.paragraph_format.line_spacing = Pt(0)
    standart2.add_run('1. Părţile contractului ', style='TNR_11').bold = True

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.paragraph_format.space_after = Pt(0)
    standart3.paragraph_format.line_spacing = Pt(0)
    standart3.add_run('  1. Universitatea de Stat din Tiraspol:', style='TNR_11').bold = True
    standart3.add_run(' adresa juridică: mun. Chişinău, str. Iablocichin 5, tel. (022) 75-49-24', style='TNR_11')



    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.paragraph_format.space_after = Pt(0)
    standart4.paragraph_format.line_spacing = Pt(0)
    standart4.add_run(
        f'  2. Studentul(a) {date_cont[0]}, Facultatea FMTI, specialitatea {date_cont[1]}, anul {date_cont[2]}, secția {date_cont[4]}, grupa {date_cont[3]}, care îşi efectueaza stagiul de practica, numit în continuare practicant. ',
        style='TNR_11')


    standart5 = doc.add_paragraph()
    standart5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart5.paragraph_format.space_after = Pt(0)
    standart5.paragraph_format.line_spacing = Pt(0)
    standart5.add_run(
        f'  3. Denumirea şi adresa juridică a instituţiei de învăţământ/organizaţiei, unde se va desfăşura stagiul de practică: {date_cont[5]}',style='TNR_11')


    standart6 = doc.add_paragraph()
    standart6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart6.paragraph_format.space_after = Pt(0)
    standart6.paragraph_format.line_spacing = Pt(0)
    standart6.add_run(f'2. Obiectul contractului ',style='TNR_11').bold = True

    standart7 = doc.add_paragraph()
    standart7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart7.paragraph_format.space_after = Pt(0)
    standart7.paragraph_format.line_spacing = Pt(0)
    standart7.add_run(f'    Prezentul contract stabileşte condiţiile în care practicantul va efectua un stagiu de practică în cadrul '
                      f'    instituţiei gazdă şi este documentul de bază în procesul de realizare a funcţiei de acumulare a creditelor ECTS.',
                      style='TNR_11')

    standart8 = doc.add_paragraph()
    standart8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart8.paragraph_format.space_after = Pt(0)
    standart8.paragraph_format.line_spacing = Pt(0)
    standart8.add_run('3. Durata contractului:', style='TNR_11').bold = True
    standart8.add_run(' de la __________ pâna la __________', style='TNR_11')

    standart9 = doc.add_paragraph()
    standart9.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart9.paragraph_format.space_after = Pt(0)
    standart9.paragraph_format.line_spacing = Pt(0)
    standart9.add_run('4. Tipul stagiului de practică: ', style='TNR_11').bold = True
    standart9.add_run(f' {date_cont[6]} ', style='TNR_11')

    standart10 = doc.add_paragraph()
    standart10.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart10.paragraph_format.space_after = Pt(0)
    standart10.paragraph_format.line_spacing = Pt(0)
    standart10.add_run('5. Locul de efectuare a practicii: ' ,style='TNR_11').bold = True
    standart10.add_run(f' {date_cont[5]} ',style='TNR_11')

    standart11 = doc.add_paragraph()
    standart11.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart11.paragraph_format.space_after = Pt(0)
    standart11.paragraph_format.line_spacing = Pt(0)
    standart11.add_run('6. Programul de efectuare a practicii: ', style='TNR_11').bold = True
    standart11.add_run(f'zilnic, de la ora {date_cont[7]} până la ora {date_cont[8]} ',style='TNR_11')

    standart12 = doc.add_paragraph()
    standart12.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart12.paragraph_format.space_after = Pt(0)
    standart12.paragraph_format.line_spacing = Pt(0)
    standart12.add_run('7. Îndrumător desemnat de către instituţia unde se desfăşoară practica' , style='TNR_11').bold = True
    standart12.add_run('(nume, prenume, funcție, telefon/adresă de contact): __________________________________________________________________',
        style='TNR_11')

    standart13 = doc.add_paragraph()
    standart13.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart13.paragraph_format.space_after = Pt(0)
    standart13.paragraph_format.line_spacing = Pt(0)
    standart13.add_run('8. Obligaţiile generale ale părţilor ', style='TNR_11').bold = True

    standart14 = doc.add_paragraph()
    standart14.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart14.paragraph_format.space_after = Pt(0)
    standart14.paragraph_format.line_spacing = Pt(0)
    standart14.add_run('    8.1. Instituţia unde se desfăşoară stagiul de practică se obligă: ', style='TNR_11').bold = True

    file = doc.add_paragraph()
    file.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    file.paragraph_format.space_after = Pt(0)
    file.paragraph_format.line_spacing = Pt(0)
    file.add_run(open('files/contract_practica1.txt', 'r', encoding='utf8').read(), style='TNR_11')

    standart15 = doc.add_paragraph()
    standart15.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart15.paragraph_format.space_after = Pt(0)
    standart15.paragraph_format.line_spacing = Pt(0)
    standart15.add_run('    8.2. Practicantului îi revin urmatoarele obligaţii: ', style='TNR_11').bold = True

    file1 = doc.add_paragraph()
    file1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    file1.paragraph_format.space_after = Pt(0)
    file1.paragraph_format.line_spacing = Pt(0)
    file1.add_run(open('files/contract_practica2.txt', 'r', encoding='utf8').read(), style='TNR_11')

    standart16 = doc.add_paragraph()
    standart16.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart16.paragraph_format.space_after = Pt(0)
    standart16.paragraph_format.line_spacing = Pt(0)
    standart16.add_run('    8.3. Universitatea de Stat din Tiraspol se obligă: ', style='TNR_11').bold = True

    file2 = doc.add_paragraph()
    file2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    file2.paragraph_format.space_after = Pt(0)
    file2.paragraph_format.line_spacing = Pt(0)
    file2.add_run(open('files/contract_practica3.txt', 'r', encoding='utf8').read(), style='TNR_11')

    standart17 = doc.add_paragraph()
    standart17.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart17.paragraph_format.space_after = Pt(0)
    standart17.paragraph_format.line_spacing = Pt(0)
    standart17.add_run('9. Evaluarea stagiilor de practică şi acumularea creditelor de studiu ', style='TNR_11').bold = True

    file3 = doc.add_paragraph()
    file3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    file3.paragraph_format.space_after = Pt(0)
    file3.paragraph_format.line_spacing = Pt(0)
    file3.add_run(open('files/contract_practica4.txt', 'r', encoding='utf8').read(), style='TNR_11')

    standart18 = doc.add_paragraph()
    standart18.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart18.paragraph_format.space_after = Pt(0)
    standart18.paragraph_format.line_spacing = Pt(0)
    standart18.add_run('Semnăturile părţilor: ',
                       style='TNR_11').bold = True

    standart19 = doc.add_paragraph()
    standart19.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart19.paragraph_format.space_after = Pt(0)
    standart19.paragraph_format.line_spacing = Pt(0)
    standart19.add_run('1. UST/decanatul FMTI               2. Instituţia/organizaţia unde se petrece practica                   3. Practicantul',
                       style='TNR_11')

    makefolder()
    doc.save(f'{path}\\Template_Contract_de_Practica.docx')
    succesMessage()

def sncs(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.39)
    section.bottom_margin = Inches(1)

    styles = doc.styles
    charstyle = styles.add_style('TNR_12', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'


    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart.paragraph_format.space_after = Pt(0)
    standart.paragraph_format.line_spacing = Inches(0.29)
    standart.add_run('UNIVERSITATEA DE STAT DIN TIRASPOL', style='TNR_12')

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.paragraph_format.space_after = Pt(0)
    standart1.paragraph_format.line_spacing = Inches(0.29)
    standart1.add_run('FACULTATEA FIZICĂ, MATEMATICĂ ȘI TEHNOLOGII INFORMAȚIONALE', style='TNR_12')

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart2.paragraph_format.space_after = Pt(0)
    standart2.paragraph_format.line_spacing = Inches(0.29)
    standart2.add_run('CONTRACT DE STUDII ÎN BAZA SNCS', style='TNR_12').underline = True

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart3.paragraph_format.space_after = Pt(0)
    standart3.paragraph_format.line_spacing = Inches(0.29)
    standart3.add_run('Nr _____din ____ septembrie 2021', style='TNR_12')

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.paragraph_format.space_after = Pt(0)
    standart4.add_run('Art. 1.', style='TNR_12')
    standart4.add_run('Părțile contractante:', style='TNR_12').bold = True

    standart5 = doc.add_paragraph()
    standart5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart5.paragraph_format.space_after = Pt(0)
    standart5.add_run(f'Facultatea FMTI, reprezentată de decan, Andrei Braicov, și studentul {date[0]} la această facultate, specialitatea {date[1]}, anul {date[2]}, grupa {date[3]}, limba de instruire română ', style='TNR_12')

    standart6 = doc.add_paragraph()
    standart6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart6.paragraph_format.space_after = Pt(0)
    standart6.add_run('Art. 2.',style='TNR_12')
    standart6.add_run('Coținutul contractului: ', style='TNR_12').bold = True
    standart6.add_run('Obligațiile facultății și ale studentului cu privire la desfășurarea activităților didactice în anul universitar 2021-2022', style='TNR_12')

    standart7 = doc.add_paragraph()
    standart7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart7.paragraph_format.space_after = Pt(0)
    standart7.add_run('Art. 3.',style='TNR_12')
    standart7.add_run('Obiectul contractului: ', style='TNR_12').bold = True
    standart7.add_run(
        'Obiectul contractului îl constituie planul individual ales de student pentru anul universitar indicat, alcătuit din disciplinele obligatorii prevăzute de planul de studii și disciplinele opționale, alese de student din oferta catedrelor. În urma alegerii făcute și stipulate în acest contract, studentul trebuie să poată acumula minimum 30 de credite / semestru și, respectiv, 60 credite / an de studii. Cantitatea de credite poate întrece limita de 60, dar nu poate depăși numărul de 80 / an de studii.',
        style='TNR_12')

    standart8 = doc.add_paragraph()
    standart8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart8.paragraph_format.space_after = Pt(0)
    standart8.add_run('     Transferarea creditelor dintr-un semestru se poate face numai prin acumularea lor în avans, în limita de 4-6 credite (suplimentar) / semestru. ', style='TNR_12')

    standart9 = doc.add_paragraph()
    standart9.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart9.paragraph_format.space_after = Pt(0)
    standart9.add_run('Art. 4.', style='TNR_12')
    standart9.add_run('Obligațiile facultății:', style='TNR_12').bold = True

    standart10 = doc.add_paragraph()
    standart10.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart10.paragraph_format.space_after = Pt(0)
    standart10.add_run('        ✓	Să pună la dispoziția studentului ', style='TNR_12')
    standart10.add_run('Regulamentul de organizare a studiilor în învățământul superior în baza Sistemului Național de Credite de Studiu (SNCS), ', style='TNR_12').italic = True
    standart10.add_run('planurile de studii, programele analitice la disciplinele ofertate;', style='TNR_12')

    standart11 = doc.add_paragraph()
    standart11.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart11.paragraph_format.space_after = Pt(0)
    standart11.add_run('        ✓	Să asigure predarea disciplinelor ofertate și alese de student prin semnarea acestui contract.', style='TNR_12')

    standart12 = doc.add_paragraph()
    standart12.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart12.paragraph_format.space_after = Pt(0)
    standart12.add_run('Art. 5.', style='TNR_12')
    standart12.add_run('Obligațiile studentului:', style='TNR_12').bold = True

    standart13 = doc.add_paragraph()
    standart13.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart13.paragraph_format.space_after = Pt(0)
    standart13.add_run('         ✓	Să respecte ', style='TNR_12')
    standart13.add_run('Regulamentul de organizare a studiilor în învățământul superior în baza Sistemului Național de Credite de Studiu (SNCS);', style='TNR_12').italic = True

    standart14 = doc.add_paragraph()
    standart14.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart14.paragraph_format.space_after = Pt(0)
    standart14.add_run('        ✓	Să urmeze disciplinele stipulate în prezentul contract, realizând toate activitățile educaționale prevăzute de programele analitice la disciplinele indicate;', style='TNR_12')

    standart15 = doc.add_paragraph()
    standart15.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart15.paragraph_format.space_after = Pt(0)
    standart15.add_run(
        '       ✓	Să susțină probele de evaluare în conformitate cu orarul stabilit de decanat.',
        style='TNR_12')

    standart16 = doc.add_paragraph()
    standart16.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart16.paragraph_format.space_after = Pt(0)
    standart16.add_run('Art. 6.', style='TNR_12')
    standart16.add_run('Disciplinele contractate: ', style='TNR_12').bold = True

    standart17 = doc.add_paragraph()
    standart17.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart17.paragraph_format.space_after = Pt(0)
    standart17.add_run(
        '       (Sunt incluse disciplinele din planul de învățământ pe care srudentul se obligă să le urmeze pe parcursul fiecărui semestru din anul universitar. În cazuri speciale, de repetare a semestrului, contractul se va perfecta pentru un semestru.)',
        style='TNR_12')

    doc.add_paragraph().add_run()

    standart18 = doc.add_paragraph()
    standart18.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart18.paragraph_format.space_after = Pt(0)
    standart18.add_run(
        'Numărul de credite angajate prin prezentul contract este:\nSemestrul 1 – 30 \nSemestrul 2 – 30 \nCertificare, metodist consilier\nSemnătura părților',
        style='TNR_12')

    doc.add_paragraph().add_run()

    standart19 = doc.add_paragraph()
    standart19.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart19.paragraph_format.space_after = Pt(0)
    standart19.add_run(
        'Decan 										Student',
        style='TNR_12')
    makefolder()
    doc.save(f'{path}\\Template_Contract_intre_Student_si_Facultate.docx')
    succesMessage()

def achitare(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    table.cell(0, 0).width = Inches(2.96)
    table.cell(0, 1).width = Inches(3.7)


    row = table.rows[0].cells
    unu = row[0].paragraphs[0]
    unu.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = unu.add_run()
    inl = r.add_picture('files/univ.png')
    inl.width = Inches(1.69)
    inl.height = Inches(1.67)

    unu2 = row[1].paragraphs[0]
    unu2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu2.add_run('APROB\n\nRector\nEduard Coropceanu\n\n_____________________', style='TNR_14')

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart.add_run('Stimate Domnule Rector,', style='TNR_14')
    

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart1.add_run(f'     Subsemnatul(a) {date[0]},', style='TNR_14')

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(f'student(ă) în anul {date[2]}, gr. {date[3]}, Facultatea {date[4]},', style='TNR_14')

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(f'specialitatea {date[1]}', style='TNR_14')

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.add_run('solicit acordul Dumneavoastră privind achitarea taxei de studii în 2 rate. Prima rată, în valoare ', style='TNR_14')

    standart5 = doc.add_paragraph()
    standart5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart5.add_run(
        f'de {date[5]}%, va fi achitată până la {date[6]}, iar a doua rată de {date[7]}% - până la ',
        style='TNR_14')

    standart6 = doc.add_paragraph()
    standart6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart6.add_run(
        f'{date[8]}, în legătură cu {date[9]}',
        style='TNR_14')

    standart7 = doc.add_paragraph()
    standart7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart7.add_run(
        'Data                                                                                Semnătura',
        style='TNR_14')

    standart8 = doc.add_paragraph()
    standart8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    standart8.add_run('\nRectorului \nUniversității de Stat din Tiraspol\nDlui Eduard Coropceanu,\ndoctor, profesor universitar     \n\n    COORDONAT\n\n   Decan', style='TNR_14')

    makefolder()
    doc.save(f'{path}\\Template_Cerere_Achitare_in_Rate.docx')
    succesMessage()

def decanat(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    table.cell(0, 0).width = Inches(2.96)
    table.cell(0, 1).width = Inches(3.7)

    row = table.rows[0].cells
    unu = row[0].paragraphs[0]
    unu.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = unu.add_run()
    inl = r.add_picture('files/dec1.jpg')
    inl.width = Inches(2.57)
    inl.height = Inches(1.54)

    unu1 = row[1].paragraphs[0]
    unu1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r1 = unu1.add_run()
    inl1 = r1.add_picture('files/dec2.jpg')
    inl1.width = Inches(2.47)
    inl1.height = Inches(1.63)

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart.add_run('Stimate Domnule Decan,', style='TNR_14')

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart1.add_run(f'Subsemnatul(a) {date[0]},', style='TNR_14')

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(f'student(ă) în anul {date[1]}, gr. {date[2]}, solicit acordul Dumneavoastră privind {date[3]}', style='TNR_14')

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(
        'Data                                                                                Semnătura',
        style='TNR_14')

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.add_run(
        '\nDecan \n\nFacultatea Fizică, Matematică \nși Tehnologii Informaționale \nAndrei Braicov\ndoctor, conferențiar universitar     ',
        style='TNR_14')

    makefolder()
    doc.save(f'{path}\\Template_Cerere_catre_Decan.docx')
    succesMessage()

def rectorat(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    table.cell(0, 0).width = Inches(2.96)
    table.cell(0, 1).width = Inches(3.7)


    row = table.rows[0].cells
    unu = row[0].paragraphs[0]
    unu.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = unu.add_run()
    inl = r.add_picture('files/univ.png')
    inl.width = Inches(1.69)
    inl.height = Inches(1.67)

    unu2 = row[1].paragraphs[0]
    unu2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    unu2.add_run('APROB\n\nRector\nEduard Coropceanu\n\n_____________________', style='TNR_14')

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart.add_run('Stimate Domnule Rector,', style='TNR_14')

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart1.add_run(f'Subsemnatul(a) {date[0]},', style='TNR_14')

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(f'student(ă) în anul {date[1]}, gr. {date[2]}, solicit acordul Dumneavoastră privind {date[3]}',
                      style='TNR_14')

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(
        'Data                                                                                Semnătura',
        style='TNR_14')

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    standart4.add_run(
        '\nRectorului \nUniversității de Stat din Tiraspol\nDlui Eduard Coropceanu,\ndoctor, profesor universitar     \n\n    COORDONAT\n\n   Decan',
        style='TNR_14')

    makefolder()
    doc.save(f'{path}\\Template_Cerere_catre_Rector.docx')
    succesMessage()

def angajare(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    standart.add_run(f'Dnei(lui) {date[2]}\n{date[3]}\n{date[1]}', style='TNR_14').bold = True

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.add_run('CERERE DE ANGAJARE', style='TNR_14').bold = True

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(
        f'    Subsemnatul (a) {date[0]}  domiciliat(ă) în {date[6]} str.{date[5]}, telefon {date[8]}.\n      Rog să-mi aprobaţi angajarea în funcţie de {date[4]},  secţia_________________________,  începînd cu data de {date[7]}.',
        style='TNR_14')
    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(
        '      Data__________                                                Semnătura__________',
        style='TNR_14')

    doc.add_paragraph().add_run()

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.add_run('      Coordonat  şef secţie:', style='TNR_14').bold = True

    makefolder()
    doc.save(f'{path}\\Template_Cerere_de_Angajare.docx')
    succesMessage()

def demisie(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    standart.add_run(f'Dnei(lui) {date[2]}\n{date[3]}\n{date[1]}', style='TNR_14').bold = True

    doc.add_paragraph().add_run()

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.add_run('CERERE DE DEMISIE', style='TNR_14').bold = True

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(
        f'    Subsemnatul (a) {date[0]}  angajat (ă) în calitate de {date[4]} secţia _______________________,   rog să–mi aprobaţi demisia din proprie iniţiativă, cu desfacerea contractului individual de muncă la data de {date[5]} (ultima zi de muncă).\n      Telefon {date[6]}'
        , style='TNR_14')

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(
        '      Data__________                                                Semnătura__________',
        style='TNR_14')

    doc.add_paragraph().add_run()

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.add_run('      Coordonat  şef secţie:', style='TNR_14').bold = True

    makefolder()
    doc.save(f'{path}\\Template_Cerere_de_Demisie.docx')
    succesMessage()

def concediu_anual(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    standart.add_run(f'Dnei(lui) {date[2]}\n{date[3]}\n{date[1]}', style='TNR_14').bold = True

    doc.add_paragraph().add_run()

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.add_run('CERERE', style='TNR_14').bold = True

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(
        f'    Subsemnatul (a) {date[0]}  angajat (ă) în funcția de {date[4]} secţia _______________________,rog să-mi acordați concediul de odihnă anual (parțial), {date[7]} zile calendaristice, începînd cu data de {date[5]}.\n     Telefon {date[6]}'
        , style='TNR_14')

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(
        '      Data__________                                                Semnătura__________',
        style='TNR_14')

    doc.add_paragraph().add_run()

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.add_run('      Coordonat  şef secţie:', style='TNR_14').bold = True

    makefolder()
    doc.save(f'{path}\\Template_Cerere_de_Concediu_Anual.docx')
    succesMessage()

def concediu_propriu(date):
    doc = docx.Document()
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.sections[0].different_first_page_header_footer = True
    sectPr = doc.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "0")
    sectPr.append(pgNumType)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.63)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.3)

    styles = doc.styles
    charstyle = styles.add_style('TNR_14', WD_STYLE_TYPE.CHARACTER)
    obj_font = charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    standart = doc.add_paragraph()
    standart.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    standart.add_run(f'Dnei(lui) {date[2]}\n{date[3]}\n{date[1]}', style='TNR_14').bold = True

    doc.add_paragraph().add_run()

    standart1 = doc.add_paragraph()
    standart1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    standart1.add_run('CERERE', style='TNR_14').bold = True

    standart2 = doc.add_paragraph()
    standart2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart2.add_run(
        f'    Subsemnatul (a) {date[0]}  angajat (ă) în funcția de {date[4]} secţia _______________________,rog să-mi acordați concediu neplătit, cu o durată de {date[7]} zile calendaristice, începînd cu data de {date[5]}, din motive personale.\n     Telefon {date[6]}'
        , style='TNR_14')

    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()

    standart3 = doc.add_paragraph()
    standart3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart3.add_run(
        '      Data__________                                                Semnătura__________',
        style='TNR_14')

    doc.add_paragraph().add_run()

    standart4 = doc.add_paragraph()
    standart4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    standart4.add_run('      Coordonat  şef secţie:', style='TNR_14').bold = True

    makefolder()
    doc.save(f'{path}\\Template_Cerere_de_Concediu_din_Cont_Propriu.docx')
    succesMessage()